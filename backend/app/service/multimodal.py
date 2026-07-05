import os
import re
import shutil
import tempfile
from typing import List, Optional,Tuple

import pdfplumber
from docx import Document as DocxDocument
from llama_index.core import Document

from ..config.config import logger
from .documents import validate_text_documents
import logging
import uuid
from pathlib import Path
from typing import List, Optional, Dict, Any
from llama_index.core import VectorStoreIndex, Settings, StorageContext
from llama_index.core.schema import BaseNode
from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding
from llama_parse import LlamaParse
from ..service.llms import build_models_from_claims


from ..service.llamparse.textprocessing import TextProcessor
from ..service.llamparse.tableprocessing import TableProcessor
from ..service.llamparse.image_extraction import extract_images_from_pdf
from ..service.llamparse.table_extraction import extract_tables_from_text
from ..service.llamparse.imageprocessing import ImageProcessor
from ..service.metadata import DocumentMetadataExtractor



class MultiModal:
    def __init__(self,llm,embed_model,original_filename,file_path : str , file_ext : str,llamaparse_api_key: str = None,standard_documents=None):
        self.llm = llm
        self.embed_model = embed_model
        self.file_path = file_path
        self.file_ext = file_ext
        self.enable_multimodal = True
        self.llamaparse_api_key = llamaparse_api_key
        self.standard_documents = standard_documents
        self.text_processor = None
        self.table_processor = None
        self.image_processor = None
        self.original_filename = original_filename
        self.use_Llamaparse = self.requires_multimodal_parsing()
        self.document_metadata: dict = {}




    def requires_multimodal_parsing(self) -> bool:
        try:
            ext = (self.file_ext or "").lower()
            if ext == ".pdf":
                with pdfplumber.open(self.file_path) as pdf:
                    for page in pdf.pages:
                        if page.extract_tables():
                            return True
                        if getattr(page, "images", None):
                            return True
                return False
            if ext == ".docx":
                doc = DocxDocument(self.file_path)
                if doc.tables:
                    return True
                try:
                    inline_shapes = doc.inline_shapes
                    if inline_shapes and len(inline_shapes) > 0:
                        return True
                except Exception:
                    pass
                return False
        except Exception as e:
            logger.warning(f"Multimodal detection failed for {self.file_path}: {e}")
            raise
        return False
    
    


    def orchestrate_document_parsing(self)->str:
        if self.enable_multimodal and self.use_Llamaparse:
            logger.info(
                f"Routing {self.file_path} through LlamaParse (tables/diagrams detected)"
            )
            try:       
                document_id = str(uuid.uuid4())
                
                logger.info(f"Processing document: {self.file_path} (document_id={document_id})")           
                logger.info("Step 1: Loading and parsing PDF document with LlamaParse...")

                if not self.llamaparse_api_key:
                    logger.error("Missing LlamaParse API key")
                    raise ValueError(
                        "This document contains tables/images and needs LlamaParse, "
                        "but no LlamaParse API key was provided at login. "
                        "Get one from https://cloud.llamaindex.ai/ and log in again with it."
                    )

                logger.info("Initializing LlamaParse parser...")

                if self.llamaparse_api_key is None:
                    raise ValueError('The LLAMA_PARSE_API_KEY is NOT FOUND')
                
                parser_lp = LlamaParse(api_key=self.llamaparse_api_key,result_type="markdown", verbose=True)
                
                logger.info(f"Parsing PDF with LlamaParse: {self.file_path}")
                documents = parser_lp.load_data(self.file_path)
                
                if not documents or not documents[0].text:
                    logger.error("No content returned by LlamaParse")
                    raise ValueError(
                        "LlamaParse returned no extractable content for this file. "
                        "The document may be empty, corrupted, or unsupported."
                    )
                
                for doc in documents:
                    metadatas = DocumentMetadataExtractor(doc=doc,llm=self.llm,file_path=self.file_path,original_filename=self.original_filename)
                    doc.metadata = metadatas.build_structural_metadata(parsed_with="llamaparse").copy()

                logger.info(f"LlamaParse returned {len(documents)} document(s)")

                full_text = "\n\n".join([doc.text for doc in documents])
                logger.info(f"Combined text length: {len(full_text)} characters")

  
                try:
                    combined_doc = Document(text=full_text)
                    self.document_metadata = DocumentMetadataExtractor(
                        doc=combined_doc,
                        llm=self.llm,
                        file_path=self.file_path,
                        original_filename=self.original_filename,
                    ).build_document_metadata(parsed_with="llamaparse")
                except Exception as meta_err:
                    logger.warning(
                        f"Document metadata enrichment failed; using structural only: {meta_err}"
                    )
                    self.document_metadata = DocumentMetadataExtractor(
                        doc=Document(text=full_text),
                        llm=self.llm,
                        file_path=self.file_path,
                        original_filename=self.original_filename,
                    ).build_structural_metadata(parsed_with="llamaparse")

                logger.info(f"Document loaded successfully, text length: {len(full_text)} characters")

                return full_text
                

            except Exception as e:
                logger.error(
                    f"LlamaParse failed for {self.file_path}; falling back to the "
                    f"LlamaIndex text load: {e}"
                )
                raise
    
    def process_docs_find(self,full_text:str,embed_model,llm) -> dict:
        all_nodes = []
        logger.info(f"Document loaded successfully, text length: {len(full_text)} characters")

        if full_text is None:
            raise ValueError('The DOCS are not loaded by Llamparse at all')
        doc_meta = getattr(self, "document_metadata", {}) or {}
        source_name = doc_meta.get("original_file_name") or doc_meta.get("file_name")

        if full_text:
            logger.info("Step 2: Processing text content with semantic chunking...")
            if self.text_processor is None:
                self.text_processor = TextProcessor(embed_model=embed_model)
            text_nodes = self.text_processor.process(
                full_text,
                metadata={
                    **doc_meta,
                    "source": source_name,
                    "modality": "text",
                    "page": "all",
                    "page_label": "all",
                    "level": "node",
                },
            )
            all_nodes.extend(text_nodes)
            logger.info(f"Created {len(text_nodes)} text nodes")
        
        logger.info("Step 3: Extracting markdown tables from document...")
        table_strings = extract_tables_from_text(full_text)
        if table_strings:
            logger.info(f"Found {len(table_strings)} table(s) in the document")
            if self.table_processor is None:
                self.table_processor = TableProcessor(llm=llm)
            for idx, table_md in enumerate(table_strings):
                logger.debug(f"Processing table {idx + 1}/{len(table_strings)}")
                table_nodes = self.table_processor.process(
                    table_md,
                    metadata={
                        **doc_meta,
                        "source": source_name,
                        "modality": "table",
                        "table_index": idx,
                        "element_label": f"table_{idx}",
                        "page_label": "all",
                        "level": "node",
                    },
                )
                all_nodes.extend(table_nodes)
                logger.debug(f"Table {idx + 1} processed, created {len(table_nodes)} node(s)")
        else:
            logger.info("No tables found in document")
        
        logger.info("Extracting images from PDF...")
        import tempfile
        with tempfile.TemporaryDirectory() as tmpdir:
            try:
                images = extract_images_from_pdf(self.file_path, tmpdir)
                logger.info(f"Extracted {len(images)} images from PDF")
                self.image_processor = ImageProcessor(llm=self.llm)
                
                for img_info in images:
                    try:
                        logger.debug(f"Processing image: {img_info['path']} (page={img_info.get('page')}, index={img_info.get('image_index')})")
                        image_nodes = self.image_processor.process_from_path(
                            img_info["path"],
                            metadata={
                                **doc_meta,
                                "source": source_name,
                                "modality": "diagram",
                                "page": img_info.get("page"),
                                "page_label": str(img_info.get("page")),
                                "image_index": img_info.get("image_index"),
                                "element_label": f"page{img_info.get('page')}_img{img_info.get('image_index')}",
                                "level": "node",
                            },
                        )
                        all_nodes.extend(image_nodes)
                        logger.debug(f"Image processed, created {len(image_nodes)} node(s) with permanent paths")
                    except Exception as e:
                        logger.error(f"Failed to process image {img_info['path']}: {e}", exc_info=True)
            except RuntimeError as e:
                logger.info(f"No images found in PDF: {e}")
        
        text_count = sum(1 for n in all_nodes if n.metadata.get('content_type') == 'text')
        table_count = sum(1 for n in all_nodes if n.metadata.get('content_type') == 'table_summary')
        image_count = sum(1 for n in all_nodes if n.metadata.get('content_type') == 'image_caption')
        
        answer = {
            "total_nodes": len(all_nodes),
            "text_nodes": text_count,
            "table_nodes": table_count,
            "image_nodes": image_count,
            "all_nodes" : all_nodes
        }
        
        logger.info(f"Total nodes created: {len(all_nodes)} (text: {text_count}, tables: {table_count}, images: {image_count})")

        return answer





_service: Optional[MultiModal] = None


def get_service() -> MultiModal:
    global _service
    if _service is None:
        _service = MultiModal()
    return _service




